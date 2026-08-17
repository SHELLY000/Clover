"""Writer for the institutional course-attainment evaluation form (Word).

The form is the artefact that is signed and archived, so it is generated
rather than transcribed: every figure in it is derived from the same
matrix that produced the workbook, and the provenance block at the foot
carries the blueprint hash, cohort hash and tool version so a reviewer
can re-derive it.
"""

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..blueprint import Blueprint
from ..model import AttainmentResult, DiagnosticReport

_BANDS = ["100-90", "89-80", "79-70", "69-60", "59-30", "<30"]
_GRADE_TEXT = {"A": "可分辨", "B": "弱可分辨", "C": "不可分辨"}


def _shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def _set(cell, text, bold=False, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return cell


def _kv_table(doc, pairs, widths=(2, 3, 2, 3)):
    t = doc.add_table(rows=len(pairs), cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, (k1, v1, k2, v2) in enumerate(pairs):
        _set(t.cell(r, 0), k1, bold=True)
        _set(t.cell(r, 1), v1)
        _set(t.cell(r, 2), k2, bold=True)
        _set(t.cell(r, 3), v2)
    return t


def write_evaluation_form(path: str, bp: Blueprint, result: AttainmentResult,
                          report: DiagnosticReport | None = None,
                          recommendations: list[str] | None = None,
                          analysis: str = "", improvement: str = "") -> str:
    c = bp.course
    pol = bp.policy
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(9)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(f"{c.institution}课程教学目标达成度评价表")
    run.bold = True
    run.font.size = Pt(15)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph().add_run("一、课程基本信息").bold = True
    _kv_table(doc, [
        ("课程编码", c.code, "课程名称", c.name),
        ("课程学分", c.credits, "课程学时", c.hours),
        ("开课学期", c.term, "年级、专业", c.cohort_label or c.program),
        ("课程性质", c.nature, "课程所属教研室", c.department),
        ("考核方式", c.assessment_mode, "达成度合格标准", f"{pol.qualifying_standard:.2f}"),
        ("任课教师", c.instructor, "参与评价学生数", f"{result.n_students} 人"),
    ])

    doc.add_paragraph().add_run("二、成绩分布").bold = True
    _rounding_note(doc, result)
    for key, label in (("final", "期末考核成绩"), ("total", "课程总评成绩")):
        dist = result.distributions.get(key)
        if not dist:
            continue
        p = doc.add_paragraph()
        p.add_run(f"{label}　平均分 {dist['mean']:.2f}　标准差 {dist['sd']:.2f}"
                  f"　有效人数 {dist['n']}").bold = True
        t = doc.add_table(rows=3, cols=len(_BANDS) + 1)
        t.style = "Table Grid"
        _set(t.cell(0, 0), "分数段", bold=True)
        _set(t.cell(1, 0), "人数", bold=True)
        _set(t.cell(2, 0), "百分比", bold=True)
        n = max(dist["n"], 1)
        for i, band in enumerate(_BANDS, start=1):
            cnt = dist["bands"].get(band, 0)
            _set(t.cell(0, i), band, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set(t.cell(1, i), cnt, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set(t.cell(2, i), f"{cnt / n * 100:.2f}%", align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().add_run("三、课程目标与毕业要求指标点达成度").bold = True
    headers = ["毕业要求指标点", "课程目标", "评价内容", "目标分值",
               "学生平均得分", "目标达成值", "95%置信区间", "指标点达成度", "识别性"]
    n_rows = 1 + sum(len(bp.items_for_objective(o.id)) for o in bp.objectives)
    t = doc.add_table(rows=n_rows, cols=len(headers))
    t.style = "Table Grid"
    for i, htxt in enumerate(headers):
        _set(t.cell(0, i), htxt, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(t.cell(0, i), "E8EEF7")

    r = 1
    for o in bp.objectives:
        items = bp.items_for_objective(o.id)
        res = next(x for x in result.objectives if x.id == o.id)
        ind = next((x for x in result.indicators if x.id == o.indicator), None)
        grade = report.grades.get(o.id, "—") if report else "—"
        start = r
        for it in items:
            share = it.allocations[o.id]
            stat = next(s for s in result.items if s.id == it.id)
            _set(t.cell(r, 2), it.name)
            _set(t.cell(r, 3), f"{share:.2f}", align=WD_ALIGN_PARAGRAPH.CENTER)
            _set(t.cell(r, 4), f"{stat.mean_ratio * share:.2f}",
                 align=WD_ALIGN_PARAGRAPH.CENTER)
            _set(t.cell(r, 5), f"{stat.mean_ratio:.2f}",
                 align=WD_ALIGN_PARAGRAPH.CENTER)
            r += 1
        _set(t.cell(start, 0), f"{o.indicator} {bp.indicator(o.indicator).text}")
        _set(t.cell(start, 1), f"{o.id}　{o.text}")
        _set(t.cell(start, 6), f"[{res.ci_low:.3f}, {res.ci_high:.3f}]",
             align=WD_ALIGN_PARAGRAPH.CENTER)
        _set(t.cell(start, 7), f"{ind.attainment:.3f}" if ind else "—",
             align=WD_ALIGN_PARAGRAPH.CENTER)
        cell = _set(t.cell(start, 8), f"{grade} {_GRADE_TEXT.get(grade, '')}",
                    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(cell, {"A": "E3F3E6", "B": "FFF4E5", "C": "FBE3E3"}.get(grade, "FFFFFF"))
        if len(items) > 1:
            for col in (0, 1, 6, 7, 8):
                t.cell(start, col).merge(t.cell(r - 1, col))

    p = doc.add_paragraph()
    p.add_run(f"课程总体达成度：{result.course_attainment:.3f}"
              f"（目标分值 {sum(o.target_points for o in result.objectives):.0f}，"
              f"平均得分 {result.course_mean_points:.2f}，"
              f"合格标准 {pol.qualifying_standard:.2f}，"
              f"{'达成' if result.course_attainment >= pol.qualifying_standard else '未达成'}）"
              ).bold = True

    if report is not None:
        doc.add_paragraph().add_run("四、评价结果的可解释性说明").bold = True
        para = doc.add_paragraph()
        para.add_run(
            f"本次评价的 {report.n_objectives} 个课程目标得分矩阵的有效秩为 "
            f"{report.effective_rank:.2f}，目标区分度指数为 "
            f"{report.separation_index:.3f}。").font.size = Pt(9)
        c_list = [k for k, v in report.grades.items() if v == "C"]
        if c_list:
            warn = doc.add_paragraph()
            run = warn.add_run(
                f"其中 {'、'.join(c_list)} 的评价证据不足以将其与共用同一证据的其他目标"
                f"区分开，上表中这些目标的达成值应作为课程整体达成度的分解结果理解，"
                f"不宜单独作为该目标已达成的证据。")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xB0, 0x2A, 0x2A)
        fails = [ck for ck in report.checks if ck.level == "fail"]
        if fails:
            doc.add_paragraph().add_run("主要诊断结论：").bold = True
            for ck in fails[:12]:
                doc.add_paragraph(f"[{ck.code}] {ck.message}", style="List Bullet")

    doc.add_paragraph().add_run("五、达成度分析与持续改进").bold = True
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    _set(t.cell(0, 0), "课程组目标达成度分析", bold=True)
    _set(t.cell(0, 1), analysis or _default_analysis(result, report))
    _set(t.cell(1, 0), "课程组持续改进意见", bold=True)
    _set(t.cell(1, 1), improvement or _default_improvement(recommendations or []))

    t = doc.add_table(rows=2, cols=4)
    t.style = "Table Grid"
    for r_, (a, b) in enumerate((("分析人签字", ""), ("教研室主任签字", ""))):
        _set(t.cell(r_, 0), a, bold=True); _set(t.cell(r_, 1), b)
    _set(t.cell(0, 2), "教学院长签字", bold=True)
    _set(t.cell(1, 2), "分析日期", bold=True)

    foot = doc.add_paragraph()
    run = foot.add_run(
        f"本表由 CLOVER {result.tool_version} 自动生成于 {result.generated_at}；"
        f"blueprint={result.blueprint_id} ({result.blueprint_hash[:12]})，"
        f"cohort={result.cohort_hash[:12]}，"
        f"自助抽样 {bp.policy.bootstrap_iterations} 次，随机种子 {bp.policy.seed}。"
        f"可用 `clover verify` 复算校验。")
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(path)
    return path


def _rounding_note(doc, result: AttainmentResult) -> None:
    """State the rounding convention, and reconcile against the other one.

    Registrar tables round each student's overall mark before tabulating;
    a blueprint total does not. The gap is small but lands on band
    boundaries, so a reviewer comparing this form against the student
    records system will see a handful of students move between bands. It
    is cheaper to explain that here than to be asked about it later.
    """
    d = result.distributions.get("total")
    alt = result.distributions.get("total_alternate_rounding")
    if not d or not alt:
        return
    label = {"none": "未取整（蓝图分值合计）", "integer": "按整数取整"}
    p = doc.add_paragraph()
    run = p.add_run(
        f"本表总评成绩口径：{label.get(d['rounding'], d['rounding'])}，"
        f"平均分 {d['mean']:.2f}。作为对照，采用"
        f"{label.get(alt['rounding'], alt['rounding'])}口径时平均分为 "
        f"{alt['mean']:.2f}，各分数段人数为 "
        f"{' / '.join(str(v) for v in alt['bands'].values() if v) or '0'}。"
        f"两者差异源自取整而非计算，教务系统的成绩册通常采用取整口径。")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _default_analysis(result: AttainmentResult,
                      report: DiagnosticReport | None) -> str:
    ok = result.course_attainment >= result.policy.qualifying_standard
    lo = min(result.objectives, key=lambda o: o.attainment)
    hi = max(result.objectives, key=lambda o: o.attainment)
    txt = (f"课程总体达成度 {result.course_attainment:.3f}，"
           f"{'高于' if ok else '低于'}合格标准 "
           f"{result.policy.qualifying_standard:.2f}。各课程目标中，"
           f"{hi.id} 最高（{hi.attainment:.3f}），"
           f"{lo.id} 最低（{lo.attainment:.3f}），极差 "
           f"{hi.attainment - lo.attainment:.3f}。")
    if report is not None and report.effective_rank < 1.5:
        txt += ("需要说明的是，各目标得分由同一批聚合成绩按固定比例分解得到，"
                "目标之间的差异反映的是分值配比而非学生在不同目标上的实际差异，"
                "因此上述极差不宜作为“某目标相对薄弱”的判断依据。")
    return txt


def _default_improvement(recs: list[str]) -> str:
    if not recs:
        return "本轮评价未发现需要立即处理的证据结构问题，建议保持现有考核方案。"
    return "　".join(f"{i}. {r}" for i, r in enumerate(recs, start=1))
